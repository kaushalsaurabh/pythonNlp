'''This program will read text from word'''
import docx


# Function to read text from dox
def get_text_from_word (filename):

    # Load the doc object with document to be read
    print('filename='+filename)
    doc = docx.Document(filename)

    # Read the doc para by para
    text = []
    for para in doc.paragraphs:
        text.append(para.text)

    return '\n'.join(text)

# Function to get multiple attributes in word file
def get_attributes(filename):
    # Create doc object
    doc = docx.Document(filename)

    # Paragraphs
    print('Number of paragraphs', len(doc.paragraphs))
    print('Second Paragraph', doc.paragraphs[1].text)
    print('Second Paragraph style', doc.paragraphs[1].style)

    # Runs. A run is continous text with same styling
    print('Paragraph 1:', doc.paragraphs[0].text)
    print('Number of runs in paragraph 1:', len(doc.paragraphs[0].runs))
    for idx, run in enumerate(doc.paragraphs[0].runs):
        print('Run %s : %s' % (idx, run.text))


print(get_attributes('samples/sample_docx_1page.docx'))
